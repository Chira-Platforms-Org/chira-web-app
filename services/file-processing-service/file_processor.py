from typing import Dict, Any, List, Optional
import os
import json
import uuid
from datetime import datetime
import pandas as pd
import PyPDF2
from docx import Document
import io

class FileProcessor:
    """
    Handles more complex file processing operations for different file types
    """
    
    def __init__(self, base_dir: str = None):
        """Initialize the processor with a base directory for processed files"""
        self.base_dir = base_dir or os.path.join(os.path.dirname(os.path.abspath(__file__)), "processed_files")
        os.makedirs(self.base_dir, exist_ok=True)
    
    async def process_document(self, file_path: str, original_filename: str) -> Dict[str, Any]:
        """Process document files like PDF, DOC, DOCX, TXT"""
        file_extension = os.path.splitext(original_filename)[1].lower()
        
        result = {
            "file_type": "document",
            "extension": file_extension,
            "processing_time": datetime.now().isoformat(),
            "statistics": {
                "file_size_bytes": os.path.getsize(file_path),
            },
            "extracted_text": ""
        }
        
        if file_extension == '.pdf':
            result["document_type"] = "PDF"
            try:
                extracted_text = []
                with open(file_path, 'rb') as pdf_file:
                    pdf_reader = PyPDF2.PdfReader(pdf_file)
                    result["statistics"]["page_count"] = len(pdf_reader.pages)
                    
                    # Extract text from each page
                    for i in range(len(pdf_reader.pages)):
                        page = pdf_reader.pages[i]
                        extracted_text.append(page.extract_text())
                    
                result["extracted_text"] = "\n\n".join(extracted_text)
                result["statistics"]["characters"] = len(result["extracted_text"])
                result["statistics"]["words"] = len(result["extracted_text"].split())
            except Exception as e:
                result["processing_error"] = f"PDF processing error: {str(e)}"
                
        elif file_extension in ['.doc', '.docx']:
            result["document_type"] = "Word Document"
            try:
                doc = Document(file_path)
                all_paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
                result["extracted_text"] = "\n".join(all_paragraphs)
                result["statistics"]["paragraph_count"] = len(doc.paragraphs)
                result["statistics"]["characters"] = len(result["extracted_text"])
                result["statistics"]["words"] = len(result["extracted_text"].split())
            except Exception as e:
                result["processing_error"] = f"Word document processing error: {str(e)}"
                
        elif file_extension == '.txt':
            result["document_type"] = "Plain Text"
            try:
                with open(file_path, 'r', errors='ignore') as text_file:
                    content = text_file.read()
                    result["extracted_text"] = content
                    result["statistics"]["word_count"] = len(content.split())
                    result["statistics"]["line_count"] = content.count('\n') + 1
                    result["statistics"]["characters"] = len(content)
            except Exception as e:
                result["processing_error"] = f"Text processing error: {str(e)}"
        
        return result
    
    async def process_tabular_data(self, file_path: str, original_filename: str) -> Dict[str, Any]:
        """Process tabular data files like CSV, XLSX"""
        file_extension = os.path.splitext(original_filename)[1].lower()
        
        result = {
            "file_type": "tabular_data",
            "extension": file_extension,
            "processing_time": datetime.now().isoformat(),
            "statistics": {
                "file_size_bytes": os.path.getsize(file_path),
            },
            "extracted_data": "",
            "summary": {}
        }
        
        try:
            if file_extension == '.csv':
                result["data_format"] = "CSV"
                # Use pandas to read and analyze the CSV
                df = pd.read_csv(file_path)
                result["statistics"]["total_rows"] = len(df)
                result["statistics"]["total_columns"] = len(df.columns)
                result["statistics"]["column_names"] = df.columns.tolist()
                result["summary"] = {
                    "column_types": {col: str(df[col].dtype) for col in df.columns},
                    "preview_rows": df.head(5).to_dict(orient="records")
                }
                # Store the data as JSON
                result["extracted_data"] = df.to_json(orient="records")
                
            elif file_extension in ['.xlsx', '.xls']:
                result["data_format"] = "Excel Spreadsheet"
                # Process Excel file
                df = pd.read_excel(file_path)
                result["statistics"]["total_rows"] = len(df)
                result["statistics"]["total_columns"] = len(df.columns)
                result["statistics"]["column_names"] = df.columns.tolist()
                result["summary"] = {
                    "column_types": {col: str(df[col].dtype) for col in df.columns},
                    "preview_rows": df.head(5).to_dict(orient="records") 
                }
                # Store the data as JSON
                result["extracted_data"] = df.to_json(orient="records")
                
                # Get sheet names for Excel files
                xls = pd.ExcelFile(file_path)
                result["statistics"]["sheet_names"] = xls.sheet_names
                result["statistics"]["sheet_count"] = len(xls.sheet_names)
        
        except Exception as e:
            result["processing_error"] = f"Tabular data processing error: {str(e)}"
            
        return result
    
    async def process_image(self, file_path: str, original_filename: str) -> Dict[str, Any]:
        """Process image files"""
        file_extension = os.path.splitext(original_filename)[1].lower()
        
        result = {
            "file_type": "image",
            "extension": file_extension,
            "processing_time": datetime.now().isoformat(),
            "statistics": {
                "file_size_bytes": os.path.getsize(file_path),
            }
        }
        
        # Image-specific processing would go here
        # For example, using Pillow to get image dimensions, format, etc.
        
        return result
    
    def get_processor_for_filetype(self, file_path: str, content_type: str, original_filename: str):
        """Return the appropriate processor based on file type"""
        file_extension = os.path.splitext(original_filename)[1].lower()
        
        # Choose processor based on file extension and content type
        if file_extension in ['.pdf', '.doc', '.docx', '.txt']:
            return self.process_document
        elif file_extension in ['.csv', '.xlsx', '.xls']:
            return self.process_tabular_data
        elif content_type.startswith('image/'):
            return self.process_image
        else:
            # Default processor that just returns basic file info
            return lambda fp, fn: {"file_type": "unknown", "extension": file_extension}
    
    async def process_file(self, file_path: str, content_type: str, original_filename: str) -> Dict[str, Any]:
        """Main entry point for processing any file"""
        processor = self.get_processor_for_filetype(file_path, content_type, original_filename)
        result = await processor(file_path, original_filename)
        
        # Add common metadata
        result.update({
            "original_filename": original_filename,
            "processing_id": uuid.uuid4().hex,
            "processed_at": datetime.now().isoformat()
        })
        
        # Read file binary data if needed for database storage
        with open(file_path, 'rb') as file:
            result["file_binary"] = file.read()
        
        return result
